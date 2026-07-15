from __future__ import annotations

import os
from pathlib import Path

import pytest

from project_copilot.ingestion import DoclingOfficeParser, ImportedFile, ProjectIndexer
from project_copilot.workspaces import WorkspaceManager


pytestmark = pytest.mark.skipif(
    os.getenv("PROJECT_COPILOT_RUN_DOCLING_INTEGRATION") != "1",
    reason="real Docling integration requires the pinned documents extra and local models",
)


def test_real_docling_pdf_and_docx_chunks_survive_restart(tmp_path: Path) -> None:
    from docx import Document as WordDocument
    from reportlab.pdfgen import canvas

    pdf_path = tmp_path / "defrost-sequence.pdf"
    pdf = canvas.Canvas(str(pdf_path))
    pdf.drawString(72, 760, "Synthetic HP-01 defrost control sequence")
    pdf.drawString(72, 735, "Maximum defrost duration is 300 seconds.")
    pdf.save()

    docx_path = tmp_path / "controls-review.docx"
    word = WordDocument()
    word.add_heading("Controls review", level=1)
    word.add_paragraph(
        "Decision D-DEF-01 requires the outdoor fan to remain off during defrost."
    )
    word.save(docx_path)

    parser = DoclingOfficeParser()
    pdf_chunks = parser.parse(pdf_path)
    docx_chunks = parser.parse(docx_path)

    assert any("300 seconds" in chunk.content for chunk in pdf_chunks)
    assert any(chunk.page == 1 for chunk in pdf_chunks)
    assert any("D-DEF-01" in chunk.content for chunk in docx_chunks)

    manager = WorkspaceManager(tmp_path / "runtime")
    workspace = manager.create_workspace(
        display_name="Docling smoke", project_id="docling-smoke"
    )
    indexer = ProjectIndexer(manager, office_parser=parser)
    imported = indexer.import_files(
        workspace.project_id,
        [
            ImportedFile(pdf_path.name, pdf_path.read_bytes(), "configuration"),
            ImportedFile(docx_path.name, docx_path.read_bytes(), "meeting"),
        ],
    )

    assert {source.status for source in imported} == {"indexed"}
    restarted = ProjectIndexer(WorkspaceManager(tmp_path / "runtime"))
    result = restarted.search(
        workspace.project_id, "What is the maximum defrost duration?"
    )
    assert result.refused is False
    assert any(citation.source == pdf_path.name for citation in result.citations)
    assert any(citation.page == 1 for citation in result.citations)
